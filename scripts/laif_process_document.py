#!/usr/bin/env python3
"""Document ingestion wrapper for LAIF structural assessment.

This runner extracts local document text, assesses the extracted text with the
existing LAIF assessment engine, and writes markdown/JSON artifacts plus an
append-only processing index. It does not perform OCR, network fetches, scoring
changes, certification changes, or governance/protected-artifact mutations.
"""

from __future__ import annotations

import argparse
import datetime as _dt
import hashlib
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, Iterable

REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from assessment_engine import assess, classify_document_type, generate_markdown_report

MIN_EXTRACTED_CHARACTERS = 20
DEFAULT_OUTPUT_DIR = "laif_outputs"
INDEX_FILE_NAME = "laif_processing_index.jsonl"

ASSESSMENT_MODES = ("external_framework", "laif_native")
SECTOR_CHOICES = (
    "auto",
    "general_ai_governance",
    "government_service_delivery",
    "departmental_ai_development",
    "procurement_vendor_governance",
    "clinical_ai",
    "employment_hr_ai",
    "education_ai",
)
EXTRACTOR_CHOICES = ("auto", "builtin", "docling", "markitdown", "python-docx", "pypdf")
BUILTIN_EXTENSIONS = {".txt", ".md", ".markdown"}


@dataclass
class ExtractionResult:
    text: str
    extractor_used: str
    extraction_confidence: str
    warnings: list[str] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)


class ExtractionError(RuntimeError):
    """Raised when local document text cannot be extracted safely."""


def utc_now_iso() -> str:
    return _dt.datetime.now(_dt.UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def safe_stem(stem: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9._-]+", "-", stem.strip())
    safe = re.sub(r"-+", "-", safe).strip(".-_")
    return safe[:120] or "document"


def normalize_text(text: str) -> str:
    return (text or "").replace("\r\n", "\n").replace("\r", "\n").strip()


def ensure_minimum_text(result: ExtractionResult, path: Path) -> ExtractionResult:
    text = normalize_text(result.text)
    if len(text) < MIN_EXTRACTED_CHARACTERS:
        raise ExtractionError(
            f"Extracted text from {path} is empty or below the safe minimum "
            f"threshold of {MIN_EXTRACTED_CHARACTERS} characters."
        )
    result.text = text
    return result


def extract_builtin(path: Path) -> ExtractionResult:
    if path.suffix.lower() not in BUILTIN_EXTENSIONS:
        raise ExtractionError(f"Built-in extractor supports only .txt, .md, and .markdown files: {path}")
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="utf-8-sig")
    return ExtractionResult(text=text, extractor_used="builtin", extraction_confidence="high")


def extract_text_fallback(path: Path) -> ExtractionResult:
    """Read UTF-8 text from a non-text extension for deterministic tests/smokes."""
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="utf-8-sig")
    return ExtractionResult(text=text, extractor_used="text-fallback", extraction_confidence="low", warnings=["non-text extension decoded as UTF-8 text"])


def extract_docling(path: Path) -> ExtractionResult:
    try:
        from docling.document_converter import DocumentConverter
    except ImportError as exc:
        raise ExtractionError("Docling extractor requested but docling is not installed.") from exc

    converter = DocumentConverter()
    converted = converter.convert(str(path))
    document = getattr(converted, "document", converted)
    if hasattr(document, "export_to_markdown"):
        text = document.export_to_markdown()
    elif hasattr(document, "export_to_text"):
        text = document.export_to_text()
    else:
        text = str(document)
    return ExtractionResult(text=text, extractor_used="docling", extraction_confidence="medium")


def extract_markitdown(path: Path) -> ExtractionResult:
    try:
        from markitdown import MarkItDown
    except ImportError as exc:
        raise ExtractionError("MarkItDown extractor requested but markitdown is not installed.") from exc

    converted = MarkItDown().convert(str(path))
    text = getattr(converted, "text_content", None) or getattr(converted, "markdown", None) or str(converted)
    return ExtractionResult(text=text, extractor_used="markitdown", extraction_confidence="medium")


def extract_python_docx(path: Path) -> ExtractionResult:
    if path.suffix.lower() != ".docx":
        raise ExtractionError("python-docx extractor supports only .docx files.")
    try:
        import docx
    except ImportError as exc:
        raise ExtractionError("python-docx extractor requested but python-docx is not installed.") from exc

    document = docx.Document(str(path))
    parts: list[str] = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return ExtractionResult(text="\n".join(parts), extractor_used="python-docx", extraction_confidence="medium")


def extract_pypdf(path: Path) -> ExtractionResult:
    if path.suffix.lower() != ".pdf":
        raise ExtractionError("pypdf extractor supports only .pdf files.")
    reader_class = None
    module_name = None
    try:
        from pypdf import PdfReader as reader_class  # type: ignore[assignment]
        module_name = "pypdf"
    except ImportError:
        try:
            from PyPDF2 import PdfReader as reader_class  # type: ignore[assignment]
            module_name = "PyPDF2"
        except ImportError as exc:
            raise ExtractionError("pypdf extractor requested but neither pypdf nor PyPDF2 is installed.") from exc

    reader = reader_class(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return ExtractionResult(text="\n".join(pages), extractor_used=module_name or "pypdf", extraction_confidence="medium")


def _attempt(name: str, func: Callable[[Path], ExtractionResult], path: Path) -> tuple[ExtractionResult | None, str | None]:
    try:
        return ensure_minimum_text(func(path), path), None
    except ExtractionError as exc:
        return None, str(exc)
    except Exception as exc:  # local parser failures should be clear, not silent
        return None, f"{name} extractor failed for {path}: {exc}"


def extract_document(path: Path, extractor: str = "auto") -> ExtractionResult:
    if not path.exists() or not path.is_file():
        raise ExtractionError(f"Input file does not exist or is not a regular file: {path}")

    suffix = path.suffix.lower()
    if extractor == "builtin":
        return ensure_minimum_text(extract_builtin(path), path)
    if extractor == "docling":
        return ensure_minimum_text(extract_docling(path), path)
    if extractor == "markitdown":
        return ensure_minimum_text(extract_markitdown(path), path)
    if extractor == "python-docx":
        return ensure_minimum_text(extract_python_docx(path), path)
    if extractor == "pypdf":
        return ensure_minimum_text(extract_pypdf(path), path)

    warnings: list[str] = []
    attempts: list[tuple[str, Callable[[Path], ExtractionResult]]] = []
    if suffix in BUILTIN_EXTENSIONS:
        attempts.append(("builtin", extract_builtin))
    attempts.extend((("docling", extract_docling), ("markitdown", extract_markitdown)))
    if suffix == ".docx":
        attempts.append(("python-docx", extract_python_docx))
        attempts.append(("text-fallback", extract_text_fallback))
    if suffix == ".pdf":
        attempts.append(("pypdf", extract_pypdf))
        attempts.append(("text-fallback", extract_text_fallback))

    for name, func in attempts:
        result, warning = _attempt(name, func, path)
        if result is not None:
            result.warnings.extend(warnings)
            return result
        if warning:
            warnings.append(warning)

    if not attempts:
        raise ExtractionError(
            f"Unsupported file type {suffix or '<none>'}; no extractor can handle it without optional packages."
        )
    attempted = "; ".join(warnings)
    if suffix not in BUILTIN_EXTENSIONS and suffix not in {".docx", ".pdf"}:
        raise ExtractionError(
            f"Unsupported file type {suffix or '<none>'}; no installed extractor could handle it. "
            f"Attempted extractors: {attempted}"
        )
    raise ExtractionError("Unable to extract document text. Attempted extractors: " + attempted)


def _term_hits(text: str, terms: Iterable[str]) -> int:
    lowered = (text or "").lower()
    return sum(lowered.count(term.lower()) for term in terms)


EU_AI_ACT_SIGNAL_TERMS = (
    "regulation laying down harmonised rules",
    "harmonised rules on artificial intelligence",
    "artificial intelligence act",
    "high-risk ai systems",
    "provider",
    "providers",
    "deployer",
    "deployers",
    "conformity assessment",
    "market surveillance",
    "official journal",
    "general-purpose ai model",
    "placing on the market",
    "post-market monitoring",
    "technical documentation",
    "serious incident reporting",
)

PUBLIC_SECTOR_POLICY_SIGNAL_TERMS = (
    "policy for the responsible use of ai in government",
    "responsible use of ai in government",
    "public servants must",
    "government agencies must",
    "agencies must disclose",
    "responsible ai use by agencies",
    "accountable official",
    "accountable officials",
    "human review",
    "disclose ai use",
    "ai use register",
    "ai use registers",
    "public sector policy",
    "digital transformation agency",
    "dta",
)


def eu_ai_act_broad_legal_signal(text: str) -> bool:
    lowered = (text or "").lower()
    hits = _term_hits(lowered, EU_AI_ACT_SIGNAL_TERMS)
    employment_hits = _term_hits(lowered, ("employment", "worker", "workers", "labour", "recruitment", "workplace rights"))
    return hits >= 2 and ("harmonised rules" in lowered or "artificial intelligence act" in lowered or "regulation laying down" in lowered or hits >= employment_hits + 2)


def public_sector_policy_signal(text: str) -> bool:
    return _term_hits(text, PUBLIC_SECTOR_POLICY_SIGNAL_TERMS) >= 2


def executive_policy_directive_signal(text: str) -> bool:
    lowered = (text or "").lower()
    return "executive order" in lowered and any(term in lowered for term in ("federal agencies", "agency heads", "secretaries", "secretary"))


def voluntary_risk_framework_signal(text: str) -> bool:
    lowered = (text or "").lower()
    return "risk management framework" in lowered and any(term in lowered for term in ("voluntary", "govern, map, measure", "non-sector-specific", "use-case agnostic"))


def sector_assurance_checklist_signal(text: str) -> bool:
    lowered = (text or "").lower()
    clinical_hits = _term_hits(lowered, ("clinical", "patient", "nhs", "dcb0129", "hazard log", "clinical safety case"))
    return ("digital technology assessment criteria" in lowered or "dtac" in lowered or clinical_hits >= 2) and clinical_hits >= 1


def broad_governance_framework_signal(text: str) -> bool:
    lowered = (text or "").lower()
    if eu_ai_act_broad_legal_signal(text):
        return True
    broad_terms = (
        "risk management framework", "voluntary framework", "non-sector-specific", "use-case agnostic",
        "harmonised rules", "artificial intelligence act", "conformity assessment", "market surveillance",
        "regulation laying down", "executive order", "federal agencies", "providers", "deployers",
        "trustworthy ai", "govern map measure manage", "govern, map, measure, and manage",
        "responsible use of ai in government", "public servants", "federal agencies",
    )
    legal_terms = ("regulation", "article", "official journal", "conformity", "provider", "deployer", "market surveillance")
    framework_hits = sum(lowered.count(term) for term in broad_terms)
    legal_hits = sum(lowered.count(term) for term in legal_terms)
    employment_hits = sum(lowered.count(term) for term in ("employment", "worker", "workers", "hiring", "candidate", "hr"))
    clinical_hits = sum(lowered.count(term) for term in ("clinical", "patient", "clinician", "nhs", "dcb0129", "hazard log"))
    if clinical_hits >= 2:
        return False
    return framework_hits >= 1 or (legal_hits >= 3 and legal_hits >= employment_hits)


def auto_sector(text: str) -> str:
    lowered = text.lower()
    doc_type = classify_document_type(text)
    if doc_type in {"binding_legal_instrument", "voluntary_risk_framework"}:
        return "general_ai_governance"
    if doc_type == "executive_policy_directive" or executive_policy_directive_signal(text):
        return "government_service_delivery" if any(term in lowered for term in ("federal agencies", "agency heads", "secretaries", "secretary")) else "general_ai_governance"
    if doc_type == "sector_assurance_checklist" or sector_assurance_checklist_signal(text):
        return "clinical_ai"
    if doc_type == "public_sector_policy" or public_sector_policy_signal(text):
        return "government_service_delivery"
    if eu_ai_act_broad_legal_signal(text) or voluntary_risk_framework_signal(text):
        return "general_ai_governance"
    if broad_governance_framework_signal(text):
        return "general_ai_governance"
    patterns: list[tuple[str, Iterable[str]]] = [
        ("clinical_ai", ("clinical", "patient", "clinician", "diagnosis", "medical", "healthcare", "safety incident", "dcb0129", "hazard log", "nhs")),
        ("procurement_vendor_governance", ("procurement", "vendor", "contract", "supplier", "service level", "audit access")),
        ("employment_hr_ai", ("employment", "hiring", "hr", "human resources", "candidate", "adverse action")),
        ("education_ai", ("education", "student", "academic", "school", "accessibility", "learning")),
        ("government_service_delivery", ("public service", "public sector", "government", "government agencies", "agencies must", "public servants", "accountable officials", "ai use register", "service delivery", "administrative review", "benefit", "caseworker")),
        ("departmental_ai_development", ("software development", "release", "pipeline", "model register", "rollback", "architecture")),
    ]
    scores = [(sum(lowered.count(term) for term in terms), sector) for sector, terms in patterns]
    best_score, best_sector = max(scores, key=lambda item: (item[0], -patterns.index((item[1], next(t for s, t in patterns if s == item[1])))))
    return best_sector if best_score > 0 else "general_ai_governance"

def resolve_assessment_mode(mode: str) -> str:
    return "laif_native_certification" if mode == "laif_native" else "external_framework"


def build_processing_metadata(
    *,
    input_path: Path,
    input_path_original: str,
    output_dir: Path,
    processed_at_utc: str,
    source_sha256: str,
    safe_output_stem: str,
    markdown_enabled: bool,
    json_enabled: bool,
    original_pending_path: str = "",
    stored_source_path: str = "",
) -> dict:
    return {
        "processed_at_utc": processed_at_utc,
        "input_path_original": input_path_original,
        "input_path": str(input_path),
        "runner_input_path": str(input_path),
        "original_pending_path": original_pending_path or input_path_original,
        "stored_source_path": stored_source_path or str(input_path),
        "input_file_name": input_path.name,
        "original_file_name": input_path.name,
        "original_file_stem": input_path.stem,
        "source_sha256": source_sha256,
        "safe_output_stem": safe_output_stem,
        "markdown_output_path": str(output_dir / f"{safe_output_stem}.laif.md") if markdown_enabled else "",
        "json_output_path": str(output_dir / f"{safe_output_stem}.laif.json") if json_enabled else "",
    }


def markdown_metadata_block(processing: dict, extraction: dict, assessment: dict) -> str:
    lines = [
        "## Document Processing Metadata",
        "",
        f"- **Processed at UTC / processed_at_utc:** {processing['processed_at_utc']}",
        f"- **Original input path:** {processing['input_path_original']}",
        f"- **Resolved input path / runner_input_path:** {processing['runner_input_path']}",
        f"- **Original pending path / original_pending_path:** {processing['original_pending_path']}",
        f"- **Stored source path / stored_source_path:** {processing['stored_source_path']}",
        f"- **Original file name:** {processing['original_file_name']}",
        f"- **Source SHA-256:** {processing['source_sha256']}",
        f"- **Extractor used:** {extraction['extractor_used']}",
        f"- **Extracted characters:** {extraction['extracted_characters']}",
        f"- **Assessment mode:** {assessment.get('assessment_mode', '')}",
        f"- **Sector profile:** {assessment.get('sector_profile', '')}",
        f"- **Safe output stem:** {processing['safe_output_stem']}",
        f"- **Markdown output path:** {processing['markdown_output_path']}",
        f"- **JSON output path:** {processing['json_output_path']}",
        "",
        "---",
        "",
    ]
    return "\n".join(lines)


def json_dump(path: Path, payload: dict) -> None:
    path.write_text(json.dumps(payload, indent=2, sort_keys=True) + "\n", encoding="utf-8")


def index_record(processing: dict, extraction: dict, assessment: dict, input_path: Path, document_name: str) -> dict:
    return {
        "processed_at_utc": processing["processed_at_utc"],
        "original_file_name": processing["original_file_name"],
        "input_path_original": processing["input_path_original"],
        "input_path": str(input_path),
        "runner_input_path": processing.get("runner_input_path", str(input_path)),
        "original_pending_path": processing.get("original_pending_path", processing.get("input_path_original")),
        "stored_source_path": processing.get("stored_source_path", str(input_path)),
        "source_sha256": processing["source_sha256"],
        "safe_output_stem": processing["safe_output_stem"],
        "markdown_output_path": processing["markdown_output_path"],
        "json_output_path": processing["json_output_path"],
        "document_name": document_name,
        "assessment_mode": assessment.get("assessment_mode"),
        "sector_profile": assessment.get("sector_profile"),
        "document_type": assessment.get("document_type"),
        "formal_laif_native_compliance": assessment.get("formal_laif_native_compliance", assessment.get("formal_laif_compliance")),
        "overall_readiness_score": assessment.get("overall_readiness_score", assessment.get("overall_score")),
        "evidence_trace_count": len(assessment.get("evidence_traces", [])),
        "remediation_patch_count": len(assessment.get("remediation_patches", [])),
        "calibration_caution_count": len(assessment.get("calibration_cautions", [])),
        "gaming_risk_note_count": len(assessment.get("gaming_risk_notes", [])),
        "extractor_used": extraction.get("extractor_used"),
        "extracted_characters": extraction.get("extracted_characters"),
        "extraction_confidence": extraction.get("extraction_confidence"),
        "warning_count": extraction.get("warning_count"),
        "error_count": extraction.get("error_count"),
    }


def append_index(output_dir: Path, record: dict) -> None:
    index_path = output_dir / INDEX_FILE_NAME
    with index_path.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(record, sort_keys=True) + "\n")



SIGNAL_CATEGORIES: list[tuple[str, tuple[str, ...], str]] = [
    ("risk management", ("risk management", "risk assessment", "risk", "govern", "map", "measure", "manage"), "source_risk_management"),
    ("human oversight", ("human oversight", "human review", "override", "intervention", "clinician review"), "operational_control"),
    ("clinical safety", ("clinical safety", "patient", "hazard log", "safety case", "DCB0129"), "sector_safety"),
    ("technical security", ("secure", "security", "resilient", "cybersecurity"), "technical_control"),
    ("data protection", ("data protection", "data governance", "data quality"), "data_governance"),
    ("privacy", ("privacy", "privacy-enhanced", "confidentiality"), "privacy_control"),
    ("audit/documentation", ("audit", "documentation", "document", "evidence", "record", "traceability"), "evidence_artifact"),
    ("monitoring/review", ("monitor", "review", "evaluate", "assessment", "supervision"), "monitoring_review"),
    ("incident reporting", ("incident", "reporting", "safety incident"), "incident_response"),
    ("accountability/owner", ("accountability", "accountable", "owner", "responsible", "assign"), "ownership"),
    ("enforcement/consequence", ("enforcement", "consequence", "penalty", "shall", "must"), "enforcement"),
    ("lifecycle/change control", ("lifecycle", "change control", "change", "release"), "lifecycle_control"),
    ("rollback/fallback", ("rollback", "fallback", "fail-safe"), "fallback"),
    ("residual risk", ("residual risk", "accepted risk", "risk acceptance"), "residual_risk"),
    ("appeal/redress/contestability", ("appeal", "redress", "contestability", "contest", "administrative review"), "redress"),
    ("procurement/supplier assurance", ("procurement", "supplier", "vendor", "contract", "assurance"), "supplier_assurance"),
    ("interoperability", ("interoperability", "interoperate", "integration"), "interoperability"),
    ("bias/discrimination/fairness", ("bias", "discrimination", "fairness", "fair", "non-discrimination"), "fairness"),
]

GAP_BLUEPRINTS = [
    ("evidence_presence_without_sufficiency", "Evidence is requested but sufficiency is not closed", "high"),
    ("obligation_without_owner", "Obligation is present without a named operational owner", "high"),
    ("risk_without_closure_gate", "Risk language lacks a closure gate", "high"),
    ("monitoring_without_threshold", "Monitoring lacks thresholds and escalation", "medium"),
    ("policy_without_enforcement_consequence", "Policy language lacks enforcement consequence", "medium"),
    ("safety_case_without_live_review", "Safety case lacks live review cadence", "high"),
    ("supplier_duty_without_deployer_acceptance", "Supplier duty lacks deployer acceptance evidence", "medium"),
    ("incident_reporting_without_redress", "Incident reporting lacks affected-person redress", "high"),
    ("lifecycle_without_change_control", "Lifecycle language lacks change-control artifact", "medium"),
    ("residual_risk_without_acceptance", "Residual risk lacks acceptance authority", "medium"),
    ("framework_guidance_without_implementation_artifact", "Framework guidance lacks implementation artifact", "high"),
    ("legal_obligation_without_operational_mapping", "Legal obligation lacks operational mapping", "high"),
]

NOISE_RE = re.compile(r"(?:[A-Za-z]{1}\s){8,}|[\ufffd]{2,}|(?:\b\w\b\s*){12,}")

STRONG_QUOTE_TERMS = (
    "shall", "must", "requires", "require", "ensure", "establish", "implement",
    "monitor", "review", "risk", "oversight", "accountability", "accountable",
    "evidence", "documentation", "document", "incident", "safety", "privacy",
    "security", "conformity", "assessment", "control", "audit", "record",
)
ACTION_QUOTE_TERMS = STRONG_QUOTE_TERMS + ("manage", "assign", "maintain", "protect", "report", "disclose")
BOILERPLATE_QUOTE_RE = re.compile(r"(difficulties with accessing|accessibility|contact us|support@|@\w|email:|telephone|copyright|isbn|all rights reserved|certain commercial entities, equipment, or materials may be identified)", re.IGNORECASE)
GENERIC_FRAGMENT_RE = re.compile(r"(?:requirements agencies must follow|monitoring, will help ensure|this document in order to describe|to combat this risk, the federal government will ensure that the collection|the assessment must be documented and take|the notification shall contain the conclusions of the assessment|by la ying down those r ules)", re.IGNORECASE)
PDF_INTR_WORD_DAMAGE_RE = re.compile(r"\b(?:super vision|inv estig ation|enf or cement|monitor ing|obliga tion|ar ticle|ser ious|general-pur pose|g eneral-pur pose|provid er|provid ed|ensur ing|har monisation|uni on|f alsifi ed|accompanie d|r isk|la ying|r ules|a rtificial|i ntelligence|p rovider|d eployer|o bligation|a ssessment|syste ms?|g enerated|cont ent|ai-g enerated|f or|exper ience|regard ing|marke t|comp et ent author ity|author ity|comp et ent|super visory|notifi cation|docu mentation|imple mentation|imple ment|assess ment|require ments|deci sions?)\b", re.IGNORECASE)
INCOMPLETE_END_RE = re.compile(r"\b(?:are|is|and|or|to|of|the|that|with|for|take|taken|should|must|shall|will|through|within|including|regarding|by|from|under|related\s+to|in\s+relation\s+to|as\s+part\s+of|in\s+order\s+to)\s*$", re.IGNORECASE)
TOC_QUOTE_RE = re.compile(r"^\s*(?:\d+(?:\.\d+)*\s+){0,2}[A-Z][A-Za-z&/ -]{2,70}\s+\d{1,4}\s*$")
TITLE_ONLY_RE = re.compile(r"^[A-Z][A-Za-z0-9&/:,() -]{8,80}$")
BROKEN_GLYPH_RE = re.compile(r"\b[A-Za-z]{1,3}(?:\s+[A-Za-z]{1,3}){4,}\b", re.IGNORECASE)


def _sentence_spans(text: str) -> list[tuple[int, int, str]]:
    spans: list[tuple[int, int, str]] = []
    for match in re.finditer(r"[^.!?\n]*(?:[.!?]|\n|$)", text):
        sent = match.group(0).strip()
        if len(sent) < 20:
            continue
        start = text.find(sent, match.start())
        if start >= 0:
            spans.append((start, start + len(sent), sent))
    if not spans and text.strip():
        snippet = text.strip()[:500]
        start = text.find(snippet)
        spans.append((start, start + len(snippet), snippet))
    return spans


DISPLAY_QUOTE_REPAIRS: tuple[tuple[str, str], ...] = (
    ("AI-g enerated", "AI-generated"),
    ("g eneral-pur pose", "general-purpose"),
    ("general-pur pose", "general-purpose"),
    ("T o", "To"),
    ("appropr iate", "appropriate"),
    ("identifie d", "identified"),
    ("pro vider", "provider"),
    ("bef ore", "before"),
    ("ser vice", "service"),
    ("repor t", "report"),
    ("managem ent", "management"),
    ("refer red", "referred"),
    ("suc h", "such"),
    ("Super vision", "Supervision"),
    ("inv estig ation", "investigation"),
    ("enf or cement", "enforcement"),
    ("monitor ing", "monitoring"),
    ("exper ience", "experience"),
    ("syste ms", "systems"),
    ("regard ing", "regarding"),
    ("marke t", "market"),
    ("comp et ent author ity", "competent authority"),
    ("author ity", "authority"),
    ("comp et ent", "competent"),
    ("super visory", "supervisory"),
    ("notifi cation", "notification"),
    ("docu mentation", "documentation"),
    ("imple mentation", "implementation"),
    ("imple ment", "implement"),
    ("assess ment", "assessment"),
    ("require ments", "requirements"),
    ("deci sions", "decisions"),
    ("deci sion", "decision"),
    ("obliga tion", "obligation"),
    ("Ar ticle", "Article"),
    ("provid er", "provider"),
    ("provid ed", "provided"),
    ("ensur ing", "ensuring"),
    ("ser ious", "serious"),
    ("syste m", "system"),
    ("har monisation", "harmonisation"),
    ("f alsifi ed", "falsified"),
    ("accompanie d", "accompanied"),
    ("cont ent", "content"),
    ("r isk", "risk"),
    ("la ying", "laying"),
    ("r ules", "rules"),
    ("a rtificial", "artificial"),
    ("i ntelligence", "intelligence"),
    ("p rovider", "provider"),
    ("d eployer", "deployer"),
    ("o bligation", "obligation"),
    ("a ssessment", "assessment"),
    ("f or", "for"),
)



SPLIT_WORD_DAMAGE_REASON = "display quote contains unresolved PDF split-word extraction damage"
SPLIT_WORD_ALLOWED_TOKENS = {
    "ai", "eu", "us", "uk", "iso", "nist", "act", "article", "risk",
    "to", "of", "in", "on", "by", "for", "and", "or", "the", "a", "an",
    "be", "its", "it", "as", "has", "have", "not", "shall", "must", "should",
    "will", "may", "can", "under", "within", "through",
}
SPLIT_WORD_COMMON_SUFFIXES = (
    "ed", "ing", "ion", "tion", "sion", "ment", "ity", "er", "or", "age",
    "ice", "ate", "ive", "al", "ent", "ant", "ure", "ary", "ence", "ance",
)
SPLIT_WORD_COMMON_PREFIXES = (
    "appropr", "identifie", "pro", "bef", "ser", "rep", "managem",
    "refer", "suc", "dam", "provid", "assess", "imple", "govern",
    "document", "monitor", "supervis", "author", "compet", "notifi",
)
GOVERNANCE_JOINED_TERMS = {
    "appropriate", "identified", "provider", "before", "service", "report",
    "management", "referred", "such", "damage", "assessment", "implementation",
    "documentation", "monitoring", "supervision", "authority", "competent",
    "notification", "requirements", "decisions", "obligation", "incident",
    "governance", "conformity", "oversight", "artificial", "intelligence",
}


def unresolved_split_word_damage(display_quote: str) -> tuple[bool, str]:
    """Detect likely unresolved PDF intra-word split damage in display text.

    The detector is intentionally heuristic and generic: it looks for unnatural
    adjacent alphabetic chunks that resemble one word after deterministic display
    repairs have already run, while allowing ordinary short governance words and
    acronyms such as AI, EU, US, UK, ISO, NIST, Act, and Article.
    """
    clean = " ".join((display_quote or "").split())
    if not clean:
        return False, ""

    suspicious: list[str] = []
    if re.search(r"(?:^|[.!?]\s+)[A-Za-z]\s+[a-z](?=\s)", clean):
        suspicious.append("sentence-start single-letter split")

    tokens = list(re.finditer(r"[A-Za-z]+", clean))
    for left, right in zip(tokens, tokens[1:]):
        if clean[left.end():right.start()] != " ":
            continue
        left_text = left.group(0)
        right_text = right.group(0)
        left_lower = left_text.lower()
        right_lower = right_text.lower()
        if left_lower in SPLIT_WORD_ALLOWED_TOKENS or right_lower in SPLIT_WORD_ALLOWED_TOKENS:
            continue
        if left_text.isupper() or right_text.isupper():
            continue
        joined = left_lower + right_lower
        if len(joined) < 6:
            continue
        suffix_like = len(right_lower) <= 5 and right_lower.endswith(SPLIT_WORD_COMMON_SUFFIXES)
        prefix_like = left_lower in SPLIT_WORD_COMMON_PREFIXES
        governance_join = joined in GOVERNANCE_JOINED_TERMS
        short_tail = len(right_lower) <= 2 and len(left_lower) >= 4
        balanced_chunks = 2 <= len(left_lower) <= 8 and 2 <= len(right_lower) <= 5 and governance_join
        if governance_join or (prefix_like and (suffix_like or len(right_lower) <= 5)) or (short_tail and prefix_like) or balanced_chunks:
            suspicious.append(f"{left_text} {right_text}")

    if suspicious:
        return True, SPLIT_WORD_DAMAGE_REASON
    return False, ""

PRIMARY_QUOTE_QUALITY_THRESHOLD = 70
FINAL_PRIMARY_GATE_PREFIX = "final primary quote admission gate failed"


def validate_primary_quote_record(record: dict) -> tuple[bool, str]:
    """Final hard admission gate for primary evidence quote records.

    This validator runs only after exact/display quote fields, display
    normalisation metadata, quote quality fields, and low-confidence metadata
    exist.  It is deliberately independent from quote-quality scoring so no
    damaged or incomplete record can leak into primary evidence rendering.
    """
    display_quote = " ".join(str(record.get("display_quote") or record.get("exact_quote") or "").split())
    exact_quote = " ".join(str(record.get("exact_quote") or "").split())
    if not display_quote:
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: empty display quote"
    if not exact_quote:
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: missing exact quote trace"
    if record.get("raw_exact_quote_retained") is not True:
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: raw exact quote trace was not retained"

    score = record.get("quote_quality_score")
    try:
        score_value = int(score)
    except (TypeError, ValueError):
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: missing quote quality score"
    if score_value < PRIMARY_QUOTE_QUALITY_THRESHOLD:
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: quote quality score below primary threshold"

    low_confidence_reason = str(record.get("low_confidence_reason") or "").strip()
    if low_confidence_reason:
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: low-confidence reason present: {low_confidence_reason}"

    unresolved_damage, unresolved_reason = unresolved_split_word_damage(display_quote)
    if unresolved_damage:
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: {unresolved_reason}"

    incomplete_reason = _incomplete_quote_reason(display_quote)
    if incomplete_reason:
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: {incomplete_reason}"

    if NOISE_RE.search(display_quote) or BROKEN_GLYPH_RE.search(display_quote):
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: obvious PDF spacing artefact remains in display quote"
    if re.search(r"\b[A-Za-z]{1,2}\s+[A-Za-z]{1,2}\s+[A-Za-z]{1,2}\b", display_quote):
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: display quote is unreadable as a normal institutional quote"

    display_letters = re.findall(r"[A-Za-z]", display_quote)
    if len(display_letters) < max(10, len(display_quote) // 4):
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: display quote lacks readable institutional text"

    lower = display_quote.lower()
    if not any(term in lower for term in ACTION_QUOTE_TERMS):
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: quote cannot support a complete source-says proposition"
    if not re.search(r"\b(shall|must|should|requires?|ensure|establish|implement|monitor|review|manage|document|assign|maintain|protect|report|disclose)\b", lower):
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: quote cannot support a complete source-says proposition"
    if len(_quote_signal_dimensions(display_quote)) < 2:
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: quote lacks enough actor/action/object/context structure for primary evidence"

    # Incident-report extraction damage has proven especially prone to clean-looking
    # but semantically brittle repairs. Keep the raw evidence visible in the
    # low-confidence trace unless the final source text was already clean.
    exact_lower = exact_quote.lower()
    if "repor t" in exact_lower and ("syste ms" in exact_lower or "marke t" in exact_lower):
        return False, f"{FINAL_PRIMARY_GATE_PREFIX}: raw incident-report evidence retains multiple PDF split-word artefacts requiring source review"

    return True, ""


def _append_final_gate_reason(record: dict, reason: str) -> dict:
    gated = dict(record)
    gated.setdefault("display_quote", gated.get("exact_quote", ""))
    gated.setdefault("quote_display_normalized", False)
    gated.setdefault("quote_display_normalization_reason", "")
    gated["raw_exact_quote_retained"] = True
    existing = str(gated.get("low_confidence_reason") or "").strip()
    if existing and reason not in existing:
        gated["low_confidence_reason"] = f"{existing}; {reason}"
    else:
        gated["low_confidence_reason"] = reason
    return gated


def _finalize_primary_quote_bank(records: list[dict], low_confidence_quote_candidates: list[dict] | None = None) -> list[dict]:
    primary: list[dict] = []
    seen_low: set[str] = {q.get("exact_quote", "") for q in (low_confidence_quote_candidates or [])}
    for record in records:
        ok, reason = validate_primary_quote_record(record)
        if ok:
            primary.append(record)
            continue
        gated = _append_final_gate_reason(record, reason)
        exact_quote = gated.get("exact_quote", "")
        if low_confidence_quote_candidates is not None and exact_quote not in seen_low:
            gated["quote_id"] = f"LQ{len(low_confidence_quote_candidates)+1:03d}"
            low_confidence_quote_candidates.append(gated)
            seen_low.add(exact_quote)
    for idx, record in enumerate(primary, start=1):
        record["quote_id"] = f"Q{idx:03d}"
    return primary


def _gate_passing_primary_quotes(quote_bank: list[dict]) -> list[dict]:
    return [quote for quote in quote_bank if validate_primary_quote_record(quote)[0]]

def _preserve_initial_case(match: re.Match[str], replacement: str) -> str:
    text = match.group(0)
    if text.startswith("AI-"):
        return replacement
    if text[:1].isupper() and replacement[:1].islower():
        return replacement[:1].upper() + replacement[1:]
    return replacement


def normalize_quote_for_display(quote: str) -> tuple[str, bool, str]:
    """Return deterministic presentation text while preserving the raw exact quote elsewhere."""
    display = quote or ""
    changed_repairs: list[str] = []
    for damaged, repaired in DISPLAY_QUOTE_REPAIRS:
        pattern = re.compile(rf"(?<![A-Za-z]){re.escape(damaged)}(?![A-Za-z])", re.IGNORECASE)
        if pattern.search(display):
            display = pattern.sub(lambda m, r=repaired: _preserve_initial_case(m, r), display)
            changed_repairs.append(f"{damaged}->{repaired}")
    normalized = display != (quote or "")
    if normalized:
        return display, True, "deterministic PDF intra-word spacing repair: " + ", ".join(changed_repairs)
    return display, False, ""


def _normalized_for_quality(quote: str) -> str:
    clean = " ".join((quote or "").split())
    return normalize_quote_for_display(clean)[0]


def _quote_display_fields(exact_quote: str) -> dict:
    display_quote, normalized, reason = normalize_quote_for_display(exact_quote)
    return {
        "display_quote": display_quote,
        "quote_display_normalized": normalized,
        "quote_display_normalization_reason": reason,
        "raw_exact_quote_retained": True,
    }


def _incomplete_quote_reason(quote: str) -> str:
    clean = " ".join((quote or "").split())
    lower = clean.lower()
    if not clean:
        return "empty quote"
    if re.match(r"^[a-z,;:]", clean) or clean.endswith(",") or INCOMPLETE_END_RE.search(clean):
        return "incomplete quote could not be expanded to complete evidence statement"
    incomplete_patterns = (
        r"^after completing the manage function, plans for prioritizing risk and regular monitoring$",
        r"^ai-generated content has undergone$",
        r"^when implementing the risk management system as provided$",
        r"shall be responsible for ensuring$",
        r"\b(has undergone|as provided|for ensuring)\s*$",
    )
    if any(re.search(pattern, lower) for pattern in incomplete_patterns):
        return "incomplete quote could not be expanded to complete evidence statement"
    return ""


def _unrepaired_extraction_damage_reason(exact_quote: str, display_quote: str) -> str:
    exact_has_damage = bool(PDF_INTR_WORD_DAMAGE_RE.search(exact_quote or "") or NOISE_RE.search(exact_quote or ""))
    display_has_damage = bool(PDF_INTR_WORD_DAMAGE_RE.search(display_quote or "") or NOISE_RE.search(display_quote or "") or BROKEN_GLYPH_RE.search(display_quote or ""))
    if exact_has_damage and display_has_damage:
        return "extraction-damaged spacing could not be safely normalised"
    return ""


def _quote_signal_dimensions(quote: str) -> set[str]:
    lower = _normalized_for_quality(quote).lower()
    dimensions: set[str] = set()
    if re.search(r"\b(provider|providers|deployer|deployers|agency|agencies|organisation|organization|department|supplier|clinician|public servants?|responsible party|accountable officials?)\b", lower):
        dimensions.add("actor")
    if re.search(r"\b(shall|must|should|required|requires?|ensure|establish|implement|maintain|monitor|review|document|report|assess|disclose|manage)\b", lower):
        dimensions.add("action")
    if re.search(r"\b(risk|evidence|documentation|oversight|accountability|incident|safety|privacy|security|conformity|assessment|register|record|review|redress|technical documentation|human review)\b", lower):
        dimensions.add("object")
    if re.search(r"\b(before deployment|post-market|incident|harm|audit|assurance|approval|monitoring|escalation|exception|deployment|review)\b", lower):
        dimensions.add("context")
    return dimensions


def _profile_quote_terms(assessment: dict) -> tuple[str, ...]:
    profile = document_profile_key(assessment) if isinstance(assessment, dict) else ""
    by_profile = {
        "nist": ("risk management framework", "govern", "map", "measure", "manage", "trustworthy ai", "valid", "reliable", "safe", "secure", "accountable", "transparent", "fair", "document risks", "monitor", "review"),
        "eu_ai_act": ("providers", "deployers", "high-risk ai systems", "risk management system", "technical documentation", "conformity assessment", "post-market monitoring", "serious incident reporting", "human oversight", "general-purpose ai model"),
        "eo_14110": ("agencies shall", "secretary", "agency", "safety", "security standards", "report", "implementation", "privacy", "civil rights", "labour", "competition", "accountability"),
        "dtac": ("clinical safety case", "hazard log", "dcb0129", "clinical safety officer", "data protection", "technical security", "interoperability", "evidence submission"),
        "australian_policy": ("agencies must", "public servants", "responsible ai use", "human review", "disclose ai use", "ai use register", "accountability record", "exception", "incident reporting"),
    }
    return by_profile.get(profile, ())


def _profile_quote_score(sentence: str, assessment: dict) -> int:
    lower = sentence.lower()
    return sum(1 for term in _profile_quote_terms(assessment) if term in lower)


def _expanded_sentence_window(text: str, start: int, end: int) -> tuple[int, int, str]:
    para_start = text.rfind("\n\n", 0, start)
    para_start = 0 if para_start < 0 else para_start + 2
    para_end = text.find("\n\n", end)
    para_end = len(text) if para_end < 0 else para_end
    sent_start = max(para_start, text.rfind(".", para_start, start) + 1, text.rfind("!", para_start, start) + 1, text.rfind("?", para_start, start) + 1)
    sent_ends = [idx for idx in (text.find(".", end), text.find("!", end), text.find("?", end)) if idx >= 0]
    sent_end = min(sent_ends) + 1 if sent_ends else para_end
    expanded = text[sent_start:sent_end].strip()
    expanded_start = text.find(expanded, sent_start) if expanded else start
    if expanded and expanded_start >= 0:
        return expanded_start, expanded_start + len(expanded), expanded
    return start, end, text[start:end].strip()


def _expand_incomplete_candidate(text: str, start: int, end: int, quote: str) -> tuple[int, int, str]:
    """Expand bridge-ending evidence only to source sentence/paragraph boundaries."""
    display_quote = _normalized_for_quality(quote)
    if not _incomplete_quote_reason(display_quote):
        return start, end, quote
    expanded_start, expanded_end, expanded_quote = _expanded_sentence_window(text, start, end)
    expanded_display = _normalized_for_quality(expanded_quote)
    if (expanded_start, expanded_end, expanded_quote) != (start, end, quote) and not _incomplete_quote_reason(expanded_display):
        return expanded_start, expanded_end, expanded_quote
    return start, end, quote

def quote_quality(quote: str, extraction: dict) -> tuple[int, str]:
    """Score candidate evidence before it can enter the primary quote bank."""
    raw_clean = " ".join((quote or "").split())
    display_clean = _normalized_for_quality(raw_clean)
    lower = display_clean.lower()
    norm_lower = lower
    if extraction.get("extraction_confidence") == "low" and extraction.get("extractor_used") != "text-fallback":
        return 20, "extractor reported low confidence"
    if not raw_clean:
        return 0, "empty quote"
    if BOILERPLATE_QUOTE_RE.search(raw_clean):
        return 10, "generic disclaimer or identification boilerplate"
    if GENERIC_FRAGMENT_RE.search(raw_clean):
        return 15, "generic incomplete fragment without governance context"
    damage_reason = _unrepaired_extraction_damage_reason(raw_clean, display_clean)
    if damage_reason:
        return 15, damage_reason
    unresolved_damage, unresolved_reason = unresolved_split_word_damage(display_clean)
    if unresolved_damage:
        return 15, unresolved_reason
    if TOC_QUOTE_RE.match(display_clean):
        return 10, "table of contents or page-number fragment"
    if NOISE_RE.search(display_clean) or BROKEN_GLYPH_RE.search(display_clean):
        return 15, "possible PDF extraction noise, glyph spacing, or malformed fragment"
    if len(re.findall(r"[A-Za-z]", display_clean)) < max(10, len(display_clean) // 4):
        return 20, "insufficient alphabetic governance content"
    if TITLE_ONLY_RE.match(display_clean) and len(display_clean.split()) <= 8 and not any(t in lower for t in ACTION_QUOTE_TERMS):
        return 20, "title or isolated heading without governance action"
    incomplete_reason = _incomplete_quote_reason(display_clean)
    if incomplete_reason:
        return 25, incomplete_reason
    dimensions = _quote_signal_dimensions(display_clean)
    if len(display_clean) < 80 and not any(t in norm_lower for t in ("shall", "must", "requires", "required", "ensure", "implement")):
        return 35, "short fragment without strong obligation or control phrase"
    if not any(t in norm_lower for t in ACTION_QUOTE_TERMS):
        return 35, "no governance action, obligation, risk, evidence, or control term"
    if not re.search(r"\b(shall|must|should|requires?|ensure|establish|implement|monitor|review|manage|document|assign|maintain|protect|report|disclose)\b", norm_lower):
        return 35, "no institutional action or obligation verb explaining a governance signal"
    if len(dimensions) < 2:
        return 45, "quote lacks enough actor/action/object/context structure for primary evidence"
    score = 70 + len(dimensions) * 5
    if len(display_clean) >= 120:
        score += 5
    if any(t in norm_lower for t in ("shall", "must", "requires", "ensure", "implement", "conformity", "incident", "evidence")):
        score += 5
    return min(score, 95), "primary evidence: complete governance action with actor/control context"


def _is_low_confidence_quote(quote: str, extraction: dict) -> tuple[bool, str]:
    score, reason = quote_quality(quote, extraction)
    return score < 70, reason


def _candidate_quote_record(quote_id: str, start: int, end: int, quote: str, processing: dict, extraction: dict, assessment: dict, category: str, repair_field: str) -> dict:
    score, reason = quote_quality(quote, extraction)
    return {
        "quote_id": quote_id,
        "source_file": processing.get("stored_source_path") or processing.get("input_path"),
        "original_file_name": processing.get("original_file_name"),
        "source_sha256": processing.get("source_sha256"),
        "document_type": assessment.get("document_type", "unknown_governance_document"),
        "sector_profile": assessment.get("sector_profile", "general_ai_governance"),
        "signal_category": category,
        "exact_quote": quote,
        **_quote_display_fields(quote),
        "surrounding_context": quote,
        "start_offset": start,
        "end_offset": end,
        "extraction_confidence": extraction.get("extraction_confidence", "unknown"),
        "why_it_matters": f"Candidate source evidence for {category} analysis.",
        "what_it_proves": f"The document contains language potentially relevant to {category}.",
        "what_it_does_not_prove": "It does not prove implementation, sufficiency, legal validity, certification, or operational adoption.",
        "linked_governance_repair_field": repair_field,
        "linked_gap_ids": [],
        "low_confidence_reason": reason if score < 70 else "",
        "quote_quality_score": score,
        "quote_quality_reason": reason,
    }


def build_quote_bank(text: str, processing: dict, extraction: dict, assessment: dict, low_confidence_quote_candidates: list[dict] | None = None) -> list[dict]:
    records: list[dict] = []
    seen: set[tuple[int, int, str]] = set()
    lowered_text = text.lower()
    spans = _sentence_spans(text)
    for category, terms, repair_field in SIGNAL_CATEGORIES:
        best: tuple[int, int, str] | None = None
        matching_spans: list[tuple[int, int, str]] = []
        for start, end, sentence in spans:
            lo = _normalized_for_quality(sentence).lower()
            if any(term.lower() in lo for term in terms):
                matching_spans.append((start, end, sentence))
        if matching_spans:
            best = max(matching_spans, key=lambda item: (_profile_quote_score(item[2], assessment), quote_quality(item[2], extraction)[0], len(item[2])))
        if best is None:
            for term in terms:
                idx = lowered_text.find(term.lower())
                if idx >= 0:
                    start = max(0, idx - 160)
                    end = min(len(text), idx + 320)
                    start, end, sentence = _expanded_sentence_window(text, start, end)
                    best = (start, end, sentence)
                    break
        if best is None:
            continue
        start, end, quote = best
        start, end, quote = _expand_incomplete_candidate(text, start, end, quote)
        if (start, end, category) in seen or quote not in text:
            continue
        score, reason = quote_quality(quote, extraction)
        if score < PRIMARY_QUOTE_QUALITY_THRESHOLD:
            # Exclude low-confidence/noisy fragments from the primary quote bank.
            continue
        seen.add((start, end, category))
        context_start = max(0, start - 160)
        context_end = min(len(text), end + 160)
        records.append({
            "quote_id": f"Q{len(records)+1:03d}",
            "source_file": processing.get("stored_source_path") or processing.get("input_path"),
            "original_file_name": processing.get("original_file_name"),
            "source_sha256": processing.get("source_sha256"),
            "document_type": assessment.get("document_type", "unknown_governance_document"),
            "sector_profile": assessment.get("sector_profile", "general_ai_governance"),
            "signal_category": category,
            "exact_quote": quote,
            **_quote_display_fields(quote),
            "surrounding_context": text[context_start:context_end].strip(),
            "start_offset": start,
            "end_offset": end,
            "extraction_confidence": extraction.get("extraction_confidence", "unknown"),
            "why_it_matters": f"This is deterministic source evidence for {category} analysis.",
            "what_it_proves": f"The document contains language relevant to {category}.",
            "what_it_does_not_prove": "It does not prove implementation, sufficiency, legal validity, certification, or operational adoption.",
            "linked_governance_repair_field": repair_field,
            "linked_gap_ids": [],
            "low_confidence_reason": "",
            "quote_quality_score": score,
            "quote_quality_reason": reason,
        })
        if len(records) >= 12:
            break
    if len(records) < 12:
        for start, end, quote in spans:
            if any(existing["exact_quote"] == quote for existing in records):
                continue
            score, reason = quote_quality(quote, extraction)
            display_fields = _quote_display_fields(quote)
            if score < PRIMARY_QUOTE_QUALITY_THRESHOLD or not display_fields["quote_display_normalized"] or _incomplete_quote_reason(display_fields["display_quote"]):
                continue
            context_start = max(0, start - 160)
            context_end = min(len(text), end + 160)
            records.append({
                "quote_id": f"Q{len(records)+1:03d}",
                "source_file": processing.get("stored_source_path") or processing.get("input_path"),
                "original_file_name": processing.get("original_file_name"),
                "source_sha256": processing.get("source_sha256"),
                "document_type": assessment.get("document_type", "unknown_governance_document"),
                "sector_profile": assessment.get("sector_profile", "general_ai_governance"),
                "signal_category": "display-normalised source evidence",
                "exact_quote": quote,
                **display_fields,
                "surrounding_context": text[context_start:context_end].strip(),
                "start_offset": start,
                "end_offset": end,
                "extraction_confidence": extraction.get("extraction_confidence", "unknown"),
                "why_it_matters": "This source evidence required deterministic display normalisation while retaining exact extracted text.",
                "what_it_proves": "The document contains a complete governance signal affected by repairable PDF intra-word spacing.",
                "what_it_does_not_prove": "It does not prove implementation, sufficiency, legal validity, certification, or operational adoption.",
                "linked_governance_repair_field": "source_excerpt",
                "linked_gap_ids": [],
                "low_confidence_reason": "",
                "quote_quality_score": score,
                "quote_quality_reason": reason,
            })
            if len(records) >= 12:
                break
    if records:
        records = sorted(records, key=lambda record: 0 if record.get("quote_display_normalized") else 1)
        for idx, record in enumerate(records, start=1):
            record["quote_id"] = f"Q{idx:03d}"
    if not records:
        for start, end, quote in spans[:1]:
            start, end, quote = _expand_incomplete_candidate(text, start, end, quote)
            score, reason = quote_quality(quote, extraction)
            if score < PRIMARY_QUOTE_QUALITY_THRESHOLD or quote not in text:
                continue
            records.append({
                "quote_id": "Q001", "source_file": processing.get("stored_source_path") or processing.get("input_path"),
                "original_file_name": processing.get("original_file_name"), "source_sha256": processing.get("source_sha256"),
                "document_type": assessment.get("document_type", "unknown_governance_document"), "sector_profile": assessment.get("sector_profile", "general_ai_governance"),
                "signal_category": "audit/documentation", "exact_quote": quote, **_quote_display_fields(quote), "surrounding_context": quote,
                "start_offset": start, "end_offset": end, "extraction_confidence": extraction.get("extraction_confidence", "unknown"),
                "why_it_matters": "Fallback exact source excerpt for reviewer orientation.",
                "what_it_proves": "The quoted text exists in the extracted source.",
                "what_it_does_not_prove": "It does not prove implementation, sufficiency, legal validity, certification, or operational adoption.",
                "linked_governance_repair_field": "source_excerpt", "linked_gap_ids": [], "low_confidence_reason": "",
                "quote_quality_score": score, "quote_quality_reason": reason,
            })
    return _finalize_primary_quote_bank(records, low_confidence_quote_candidates)


def build_low_confidence_quote_candidates(text: str, processing: dict, extraction: dict, assessment: dict, limit: int = 12) -> list[dict]:
    candidates: list[dict] = []
    seen_quotes: set[str] = set()
    for start, end, quote in _sentence_spans(text):
        score, reason = quote_quality(quote, extraction)
        compact = " ".join(quote.split())
        if score >= PRIMARY_QUOTE_QUALITY_THRESHOLD or compact in seen_quotes:
            continue
        display_compact = _normalized_for_quality(compact).lower()
        if not any(term in display_compact for term in ACTION_QUOTE_TERMS + ("artificial intelligence", "ai-generated", "secure", "resilient")) and score > 15:
            continue
        seen_quotes.add(compact)
        candidates.append(_candidate_quote_record(f"LQ{len(candidates)+1:03d}", start, end, quote, processing, extraction, assessment, "low-confidence candidate", "source_excerpt"))
        if len(candidates) >= limit:
            break
    return candidates



def document_profile_key(assessment: dict) -> str:
    if assessment.get("document_profile_key"):
        return assessment["document_profile_key"]
    doc_type = assessment.get("document_type", "unknown_governance_document")
    sector = assessment.get("sector_profile", "general_ai_governance")
    text_force = str(assessment.get("governance_force_profile") or assessment.get("governance_force_summary") or "").lower()
    if sector == "clinical_ai" or doc_type == "sector_assurance_checklist":
        return "dtac"
    if doc_type == "binding_legal_instrument":
        return "eu_ai_act"
    if doc_type == "executive_policy_directive":
        return "eo_14110"
    if doc_type == "public_sector_policy" or sector == "government_service_delivery":
        return "australian_policy"
    if doc_type == "voluntary_risk_framework" or "risk-management framework" in text_force or "risk management framework" in text_force:
        return "nist"
    return doc_type


def document_specific_gap_title(gap_type: str, fallback: str, assessment: dict) -> str:
    profile = document_profile_key(assessment)
    titles = {
        "nist": {
            "framework_guidance_without_implementation_artifact": "NIST guidance requires implementation artifact before assurance reliance",
            "risk_without_closure_gate": "NIST risk guidance requires a risk closure gate before operational reliance",
            "evidence_presence_without_sufficiency": "NIST trustworthiness evidence requires sufficiency criteria before assurance reliance",
        },
        "eu_ai_act": {
            "legal_obligation_without_operational_mapping": "EU AI Act legal obligation requires local provider/deployer evidence mapping",
            "evidence_presence_without_sufficiency": "EU AI Act technical documentation requires evidence sufficiency gate",
            "monitoring_without_threshold": "EU AI Act monitoring obligation requires local threshold and escalation workflow",
        },
        "eo_14110": {
            "obligation_without_owner": "EO 14110 agency direction requires accountable implementation owner",
            "monitoring_without_threshold": "EO 14110 reporting expectation requires implementation tracking threshold",
            "policy_without_enforcement_consequence": "EO 14110 agency policy force requires escalation consequence",
        },
        "dtac": {
            "evidence_presence_without_sufficiency": "DTAC evidence submission requires sufficiency and live review gate",
            "safety_case_without_live_review": "DTAC clinical safety case requires live review and hazard-log ownership",
            "risk_without_closure_gate": "DTAC transferred clinical risk requires acceptance gate",
        },
        "australian_policy": {
            "obligation_without_owner": "Public sector AI policy requires accountable owner for responsible use",
            "evidence_presence_without_sufficiency": "Public sector AI use records require evidence sufficiency and disclosure control",
            "policy_without_enforcement_consequence": "Government AI policy requires exception and incident escalation consequence",
        },
    }
    return titles.get(profile, {}).get(gap_type, fallback)


def control_name_for_gap(gap: dict) -> str:
    profile = document_profile_key(gap)
    names = {
        "nist": ["AI Risk Management Implementation Register", "Trustworthiness Evidence Acceptance Matrix", "AI Risk Monitoring and Review Gate"],
        "eu_ai_act": ["Provider/Deployer Obligation Mapping Register", "High-Risk AI Evidence and Technical Documentation Gate", "Post-Market Monitoring and Incident Escalation Control"],
        "eo_14110": ["Agency AI Directive Implementation Tracker", "Federal AI Safety and Security Evidence Register", "Agency Accountability and Reporting Gate"],
        "dtac": ["Clinical Safety Live Assurance Register", "DTAC Evidence Sufficiency Matrix", "Transferred Clinical Risk Acceptance Register"],
        "australian_policy": ["Public Sector AI Use Register", "Human Review and Accountability Evidence Log", "AI Use Disclosure and Exception Register", "Agency AI Incident and Exception Register", "Public Sector AI Monitoring and Assurance Gate"],
    }
    idx = max(0, int(gap.get("gap_id", "GAP-001").split("-")[-1]) - 1)
    options = names.get(profile)
    if options:
        return options[idx % len(options)]
    return f"Operational closure control for {gap['gap_type'].replace('_', ' ')}"


def executive_thesis(assessment: dict, gaps: list[dict], controls: list[dict]) -> str:
    doc_type = assessment.get("document_type", "unknown_governance_document")
    force = assessment.get("governance_force_profile") or assessment.get("governance_force_summary") or "governance force requires reviewer confirmation"
    profile = document_profile_key(assessment)
    base = {
        "voluntary_risk_framework": "This document is valuable as a governance design framework, but it does not itself create binding implementation gates, evidence acceptance rules, or operational enforcement consequences.",
        "binding_legal_instrument": "The EU AI Act is a high-force legal source, but it still requires local provider/deployer obligation mapping, evidence registers, and post-market monitoring workflows before it becomes operational assurance.",
        "executive_policy_directive": "This document creates executive direction and agency expectations, but systemic repair depends on agency implementation tracking, ownership, evidence collection, and escalation consequences.",
        "sector_assurance_checklist": "This document is useful as a sector assurance screen, but it must be tied to live review, evidence sufficiency, residual risk acceptance, and deployment consequences.",
        "public_sector_policy": "The Australian Government AI policy is a public-sector operating policy. Its value depends on whether agencies maintain AI use registers, disclosure evidence, human review logs, exception records, and incident/escalation pathways.",
        "implementation_guide": "This document is useful as implementation guidance, but it must be converted into accountable owners, mandatory artifacts, thresholds, review cadence, and stop/go consequences.",
        "internal_policy": "This document is useful as institutional operating policy, but assurance depends on implementation records, owner accountability, monitoring, exceptions, incidents, and escalation evidence.",
    }.get(doc_type, "This document is useful as a governance source, but institutional reliance depends on proof of operational closure, accountable ownership, evidence sufficiency, and decision consequences.")
    area = (assessment.get("strengths") or ["governance signal requires reviewer confirmation"])[0]
    gap = gaps[0]["gap_title"] if gaps else "operational closure requires reviewer confirmation"
    action = controls[0]["control_name"] if controls else "create an accountable implementation register"
    return f"{base} Classified as `{doc_type}` with governance force `{force}`. Strongest detected control area: {area}. Principal operational gap: {gap}. Recommended next action: implement {action}."

def build_governance_gap_register(assessment: dict, quote_bank: list[dict]) -> list[dict]:
    quote_ids = [q["quote_id"] for q in quote_bank[:3]]
    scores = {k: assessment.get(k) for k in ("structural_score", "terminology_score", "conceptual_proximity_score", "auditability_score", "enforceability_score", "overall_readiness_score")}
    doc_type = assessment.get("document_type", "unknown_governance_document")
    gaps: list[dict] = []
    for idx, (gap_type, title, severity) in enumerate(GAP_BLUEPRINTS[:6], 1):
        if gap_type == "safety_case_without_live_review" and assessment.get("sector_profile") != "clinical_ai":
            continue
        if gap_type == "supplier_duty_without_deployer_acceptance" and assessment.get("sector_profile") != "procurement_vendor_governance":
            continue
        gap_id = f"GAP-{idx:03d}"
        gaps.append({
            "gap_id": gap_id,
            "gap_title": document_specific_gap_title(gap_type, title, assessment),
            "severity": severity,
            "gap_type": gap_type,
            "document_type": doc_type,
            "sector_profile": assessment.get("sector_profile"),
            "document_profile_key": document_profile_key(assessment),
            "source_evidence_quote_ids": quote_ids,
            "related_scores": scores,
            "related_governance_repair_fields": ["operational_closure", "evidence_sufficiency", "governance_force"],
            "operational_meaning": "The document may create a governance expectation, but an institution still needs owner, artifact, threshold, cadence, and decision consequence evidence.",
            "failure_mode": "Paper compliance without live operational control.",
            "affected_stakeholders": assessment.get("sector_relevant_interests", [])[:3] or ["affected people", "operators", "assurance reviewers"],
            "required_control_ids": [f"CTRL-{idx:03d}"],
            "reviewer_note": "Confirm whether implementation artifacts outside this source document close the gap.",
        })
    if not gaps:
        gaps.append({
            "gap_id": "GAP-001", "gap_title": document_specific_gap_title("framework_guidance_without_implementation_artifact", "Framework guidance lacks implementation artifact", assessment), "severity": "high",
            "gap_type": "framework_guidance_without_implementation_artifact", "document_type": doc_type,
            "sector_profile": assessment.get("sector_profile"),
            "document_profile_key": document_profile_key(assessment),
            "source_evidence_quote_ids": quote_ids, "related_scores": scores,
            "related_governance_repair_fields": ["implementation_artifact"],
            "operational_meaning": "Source language must be translated into an auditable operating control.",
            "failure_mode": "Guidance is cited as assurance without proof of adoption.",
            "affected_stakeholders": ["affected people", "assurance reviewers"], "required_control_ids": ["CTRL-001"],
            "reviewer_note": "Require local evidence before relying on this document as an assurance mechanism.",
        })
    for quote in quote_bank:
        quote["linked_gap_ids"] = [gap["gap_id"] for gap in gaps if quote["quote_id"] in gap["source_evidence_quote_ids"]]
    return gaps


def build_failure_pathways(gaps: list[dict], quote_bank: list[dict]) -> list[dict]:
    pathways: list[dict] = []
    for idx, gap in enumerate(gaps[:3], 1):
        ctrl = gap.get("required_control_ids", [f"CTRL-{idx:03d}"])[0]
        pathways.append({
            "pathway_id": f"PATH-{idx:03d}",
            "title": f"{gap['gap_title']} failure pathway",
            "severity": gap.get("severity", "medium"),
            "steps": [
                "A source document states or implies a governance expectation.",
                "The institution treats the expectation as assurance evidence.",
                "No operational owner, threshold, evidence artifact, or decision consequence is verified.",
                "An AI-enabled decision or deployment proceeds under paperwork compliance.",
                "Harm, unfairness, security exposure, or assurance failure can occur without a reliable audit trail.",
            ],
            "source_evidence_quote_ids": gap.get("source_evidence_quote_ids", []),
            "triggering_gap_ids": [gap["gap_id"]],
            "likely_institutional_failure": "Paperwork compliance without live operational control.",
            "consequence": "The reviewer cannot prove that the stated governance expectation controlled the real decision pathway.",
            "required_controls": [ctrl],
            "detection_signal": "Missing or stale owner sign-off, control evidence, threshold breach log, or review record.",
            "escalation_gate": "Pause deployment or reliance until the required artifact and accountable owner are confirmed.",
        })
    return pathways


def build_control_recommendations(gaps: list[dict], pathways: list[dict], quote_bank: list[dict]) -> list[dict]:
    controls: list[dict] = []
    path_by_gap = {gap_id: p["pathway_id"] for p in pathways for gap_id in p.get("triggering_gap_ids", [])}
    for idx, gap in enumerate(gaps, 1):
        control_id = gap.get("required_control_ids", [f"CTRL-{idx:03d}"])[0]
        controls.append({
            "control_id": control_id,
            "control_name": control_name_for_gap(gap),
            "priority": "immediate" if gap.get("severity") == "high" else "near_term",
            "risk_addressed": gap.get("failure_mode"),
            "source_evidence_quote_ids": gap.get("source_evidence_quote_ids", []),
            "linked_gap_ids": [gap["gap_id"]],
            "linked_failure_pathways": [path_by_gap.get(gap["gap_id"])] if path_by_gap.get(gap["gap_id"]) else [],
            "owner": "Named accountable business, clinical, procurement, legal, security, or assurance owner for the controlled decision pathway.",
            "required_artifact": "Signed control record linking source requirement, local procedure, evidence file, threshold, reviewer, and decision outcome.",
            "minimum_evidence": "Current owner sign-off, implementation artifact, threshold log, review cadence record, and exception/escalation register.",
            "implementation_steps": [
                "Map the quoted source expectation to a local operational requirement.",
                "Assign a named owner and independent reviewer.",
                "Define trigger, threshold, evidence artifact, cadence, and stop/go consequence.",
                "Record review outcomes and exceptions in an auditable register.",
            ],
            "trigger": "New deployment, material model/process change, supplier update, incident, threshold breach, or scheduled assurance review.",
            "threshold": "No deployment or continued reliance when required evidence is absent, stale, contradicted, or owner-unapproved.",
            "cadence": "Before deployment, after material change, after incident, and at least quarterly while in operational use.",
            "decision_consequence": "Proceed only with complete evidence; otherwise pause, remediate, escalate, or reject supplier/system use.",
            "residual_risk_if_not_implemented": "The institution may rely on governance language that does not control real-world decisions or harms.",
            "suggested_template_row": f"{control_id} | owner | artifact | trigger | threshold | cadence | decision consequence | quote IDs {', '.join(gap.get('source_evidence_quote_ids', []))}",
        })
    return controls


def _md_list(items: list[str], empty: str = "Reviewer confirmation required.") -> str:
    return "\n".join(f"- {item}" for item in items) if items else f"- {empty}"


def build_institutional_report(processing: dict, extraction: dict, assessment: dict, quote_bank: list[dict], gaps: list[dict], pathways: list[dict], controls: list[dict]) -> str:
    mode = assessment.get("assessment_mode")
    doc_type = assessment.get("document_type", "unknown_governance_document")
    force = assessment.get("governance_force_profile") or assessment.get("governance_force_summary") or "source governance force requires reviewer confirmation"
    lines = [
        f"# Institutional Governance Assessment — {processing.get('original_file_name')}", "",
        "## Executive finding", "",
    ]
    if mode == "external_framework":
        lines.append(executive_thesis(assessment, gaps, controls))
    else:
        lines.append("This document is assessed in LAIF-native mode. Formal LAIF-native certification remains governed by the deterministic LAIF validation boundary shown in the technical appendix.")
    lines += ["", "## Document identity and document type", "", f"- **Original file:** {processing.get('original_file_name')}", f"- **Document type:** {doc_type}", f"- **Assessment mode:** {mode}", f"- **Sector profile:** {assessment.get('sector_profile_label', assessment.get('sector_profile'))}", f"- **Source SHA-256:** {processing.get('source_sha256')}", "", "## Recommended use / not sufficient for", "", "- **Recommended use:** source framework review, procurement/legal/clinical/public-sector assurance scoping, control mapping, and remediation planning.", "- **Not sufficient for:** standalone proof of implementation, legal validity, external certification, supplier acceptance, clinical safety approval, or LAIF-native certification unless separately evidenced.", "", "## Governance force profile", "", f"- {force if isinstance(force, str) else json.dumps(force, sort_keys=True)}", "- The document creates a strong evidence request where it uses risk, oversight, evidence, review, incident, or accountability language, but the reviewer must test whether that request is operationally closed.", "", "## Key quoted evidence", ""]
    primary_quotes = _gate_passing_primary_quotes(quote_bank)
    for q in primary_quotes[:6]:
        quote_text = q.get("display_quote") or q["exact_quote"]
        lines.append(f"- **{q['quote_id']} — {q['signal_category']}:** “{quote_text}”")
    if not primary_quotes:
        lines.append("- No high-confidence primary quotes were extracted; use the technical appendix and source text review before relying on evidence.")
    lines += ["", "## What the document controls well", "", _md_list(assessment.get("strengths", [])[:8], "No deterministic strengths detected."), "", "## What the document does not control", ""]
    lines.append(_md_list([g["gap_title"] + f" ({g['gap_id']})" for g in gaps[:8]]))
    lines += ["", "## Hidden failure pathways", "", "Failure pathway summaries below show how paperwork compliance can proceed without live operational control.", ""]
    for pth in pathways:
        lines.append(f"### {pth['pathway_id']} — {pth['title']}")
        lines.append("")
        lines.append(_md_list(pth.get("steps", [])))
        lines.append(f"- **Escalation gate:** {pth.get('escalation_gate')}")
        lines.append("")
    lines += ["## Operational gap analysis", ""]
    for gap in gaps:
        lines.append(f"- **{gap['gap_id']} ({gap['severity']}):** {gap['operational_meaning']} Evidence: {', '.join(gap.get('source_evidence_quote_ids', [])) or 'review required'}.")
    lines += ["", "## Priority remediation roadmap", ""]
    for ctrl in controls[:8]:
        lines.append(f"- **{ctrl['priority']} — {ctrl['control_id']}:** {ctrl['control_name']}; owner: {ctrl['owner']}; artifact: {ctrl['required_artifact']}")
    lines += ["", "## Control implementation templates", "", "| Control ID | Owner | Required artifact | Trigger | Threshold | Cadence | Decision consequence |", "| --- | --- | --- | --- | --- | --- | --- |"]
    for ctrl in controls[:8]:
        lines.append(f"| {ctrl['control_id']} | {ctrl['owner']} | {ctrl['required_artifact']} | {ctrl['trigger']} | {ctrl['threshold']} | {ctrl['cadence']} | {ctrl['decision_consequence']} |")
    lines += ["", "## Residual risk if no action is taken", "", "The failure pathway is paperwork compliance without live operational control: governance language may be cited while real decisions proceed without verified owner authority, implementation artifacts, thresholds, escalation gates, or affected-person redress.", "", "## Technical appendix pointer", "", f"See `{processing.get('safe_output_stem')}.technical_appendix.md` for processing metadata, source identity, scoring table, evidence traces, remediation patches, LAIF-native construct coverage, and certification boundary.", ""]
    return "\n".join(lines)


def build_technical_appendix(processing: dict, extraction: dict, assessment: dict, quote_bank: list[dict], gaps: list[dict], pathways: list[dict], controls: list[dict], low_confidence_quote_candidates: list[dict] | None = None) -> str:
    scores = ["structural_score", "terminology_score", "conceptual_proximity_score", "auditability_score", "enforceability_score", "overall_readiness_score"]
    lines = [f"# Technical Appendix — {processing.get('original_file_name')}", "", "## Document metadata", ""]
    for key in ("original_file_name", "source_sha256", "safe_output_stem"):
        lines.append(f"- **{key}:** {processing.get(key)}")
    lines += ["", "## Processing metadata", ""]
    for key in ("processed_at_utc", "input_path_original", "original_pending_path", "stored_source_path", "runner_input_path", "markdown_output_path", "json_output_path"):
        lines.append(f"- **{key}:** {processing.get(key)}")
    lines += ["", "## Extraction metadata", ""]
    for key in ("extractor_requested", "extractor_used", "extraction_confidence", "extracted_characters", "warning_count", "warnings", "error_count", "errors", "network_access_used"):
        lines.append(f"- **{key}:** {extraction.get(key)}")
    lines += ["", "## Scoring table", "", "| Score | Value |", "| --- | --- |"]
    for key in scores:
        lines.append(f"| {key} | {assessment.get(key)} |")
    lines += ["", "## Governance repair fields", "", "```json", json.dumps({k: assessment.get(k) for k in assessment if k.startswith('governance_') or k in ('document_type','assessment_mode')}, indent=2, sort_keys=True), "```", "", "## Evidence traces", "", "```json", json.dumps(assessment.get("evidence_traces", []), indent=2, sort_keys=True), "```", "", "## Remediation patches", "", "```json", json.dumps(assessment.get("remediation_patches", []), indent=2, sort_keys=True), "```", "", "## LAIF-native construct coverage", "", "```json", json.dumps(assessment.get("construct_coverage", {}), indent=2, sort_keys=True), "```", "", "## Formal LAIF-native certification boundary", ""]
    if assessment.get("assessment_mode") == "external_framework":
        lines.append("Formal LAIF-native certification: Not claimed / not applicable to this external-framework assessment. Construct coverage is internal diagnostic data only.")
    else:
        lines.append(f"Formal LAIF-native certification: {assessment.get('formal_laif_native_compliance', assessment.get('formal_laif_compliance'))}")
    normalized_quote_ids = [q.get("quote_id") for q in quote_bank if q.get("quote_display_normalized")]
    lines += ["", "## Quote display traceability", ""]
    if normalized_quote_ids:
        lines.append("Display quote deterministically normalised from exact extracted substring; raw exact quote retained in analyst bundle.")
        lines.append(f"- Display-normalised quote IDs: {', '.join(normalized_quote_ids)}")
    else:
        lines.append("No primary quote display normalization was required; display quotes match exact extracted substrings.")
    lines += ["", "## Low-confidence extraction/noise findings", "", f"- Low-confidence extraction noise: {assessment.get('low_confidence_extraction_noise', {})}", f"- Runner warnings: {extraction.get('warnings', [])}", "", "## Low-confidence quote candidates", "", "```json", json.dumps(low_confidence_quote_candidates or [], indent=2, sort_keys=True), "```", "", "## Warnings/errors", "", f"- Warnings: {extraction.get('warnings', [])}", f"- Errors: {extraction.get('errors', [])}", ""]
    return "\n".join(lines)


def build_ai_prompt() -> str:
    return """# AI Analyst Prompt

Use only the provided deterministic analyst bundle, high-quality `quote_bank`, diagnostics, and source excerpts. Do not invent quotes, obligations, legal claims, scores, documents, actors, controls, certifications, or legal-validity conclusions.

## Required analyst approach

- Lead with a document-specific thesis that names the document type, governance force, strongest control area, principal operational gap, and next action.
- Use quote IDs only from the high-quality `quote_bank` as primary support.
- Use `display_quote` for prose quotations in the executive narrative, and cite the quote ID. Preserve `exact_quote` as the audit trace to the extracted source substring.
- Do not cite `low_confidence_quote_candidates` as primary evidence; mention them only in a technical caveat if useful.
- Do not cite incomplete fragments, heading-only quotes, boilerplate, or extraction-damaged fragments.
- If available quotes are weak, state that the deterministic quote bank is insufficient and request better source extraction.
- Use a document-specific thesis, not generic LAIF language.
- Do not invent missing source context to rescue weak quotes.
- Distinguish “source says X” from “institution has implemented X.”
- Convert generic controls into client-ready implementation actions with owner, artifact, threshold, cadence, and decision consequence.
- Explain failure pathways in plain institutional terms: what breaks, who owns it, what evidence is missing, and what decision should stop.
- Distinguish legal force from operational force. A legal or policy source may be authoritative while still lacking local implementation evidence.
- Distinguish evidence presence from evidence sufficiency. A requested document, record, or quote is not proof that evidence is current, complete, reviewed, or accepted.
- Include an executive version and a technical appendix.
- Do not treat LAIF-native failure as the headline for external-framework documents.
- Preserve all source references and quote IDs. For each major recommendation, link to quote IDs, gap IDs, and control IDs.
- If evidence is insufficient, say so. Do not claim legal validity/invalidity. Do not claim certification unless provided by deterministic LAIF data.
"""


def build_validation_rules() -> str:
    return """# AI Report Validation Rules\n\n- Every quote must exist in `quote_bank`.\n- Every recommendation must map to a gap/control ID.\n- No invented legal obligations.\n- No invented citation.\n- No unsupported certification or legal-validity claim.\n- Required sections must be present.\n- Technical appendix must be preserved.\n- Low-confidence evidence must not be used as primary support.\n"""


def write_institutional_outputs(output_dir: Path, processing: dict, extraction: dict, assessment: dict, extracted_text: str) -> dict:
    low_confidence_quote_candidates = build_low_confidence_quote_candidates(extracted_text, processing, extraction, assessment)
    quote_bank = build_quote_bank(extracted_text, processing, extraction, assessment, low_confidence_quote_candidates)
    quote_bank = _finalize_primary_quote_bank(quote_bank, low_confidence_quote_candidates)
    gaps = build_governance_gap_register(assessment, quote_bank)
    pathways = build_failure_pathways(gaps, quote_bank)
    controls = build_control_recommendations(gaps, pathways, quote_bank)
    bundle = {
        "document_metadata": {"original_file_name": processing.get("original_file_name"), "source_sha256": processing.get("source_sha256"), "document_type": assessment.get("document_type"), "sector_profile": assessment.get("sector_profile")},
        "processing_metadata": processing,
        "extraction_metadata": extraction,
        "governance_repair_fields": {k: assessment.get(k) for k in assessment if k.startswith("governance_")},
        "scores": {k: assessment.get(k) for k in ("structural_score", "terminology_score", "conceptual_proximity_score", "auditability_score", "enforceability_score", "overall_readiness_score")},
        "quote_bank": quote_bank,
        "gap_register": gaps,
        "failure_pathways": pathways,
        "control_recommendations": controls,
        "extraction_warnings": extraction.get("warnings", []),
        "low_confidence_evidence_flags": [q for q in quote_bank if q.get("low_confidence_reason")],
        "low_confidence_quote_candidates": low_confidence_quote_candidates,
        "technical_appendix_data": {"construct_coverage": assessment.get("construct_coverage", {}), "formal_laif_native_compliance": assessment.get("formal_laif_native_compliance", assessment.get("formal_laif_compliance")), "evidence_traces": assessment.get("evidence_traces", []), "remediation_patches": assessment.get("remediation_patches", [])},
    }
    analyst_dir = output_dir / "analyst"
    analyst_dir.mkdir(parents=True, exist_ok=True)
    stem = processing["safe_output_stem"]
    (output_dir / f"{stem}.institutional_report.md").write_text(build_institutional_report(processing, extraction, assessment, quote_bank, gaps, pathways, controls), encoding="utf-8")
    (output_dir / f"{stem}.technical_appendix.md").write_text(build_technical_appendix(processing, extraction, assessment, quote_bank, gaps, pathways, controls, low_confidence_quote_candidates), encoding="utf-8")
    json_dump(analyst_dir / "analyst_bundle.json", bundle)
    with (analyst_dir / "quote_bank.jsonl").open("w", encoding="utf-8") as handle:
        for quote in quote_bank:
            handle.write(json.dumps(quote, sort_keys=True) + "\n")
    quote_md = ["# Quote Bank", ""]
    for q in _gate_passing_primary_quotes(quote_bank):
        display_quote = q.get("display_quote") or q["exact_quote"]
        quote_md += [
            f"## {q['quote_id']} — {q['signal_category']}",
            "",
            f"- **Signal category:** {q['signal_category']}",
            f"- **Quality score:** {q.get('quote_quality_score')} — {q.get('quote_quality_reason')}",
            "",
            "- **Display quote:**",
            f"> {display_quote}",
            "",
        ]
        if display_quote != q["exact_quote"]:
            quote_md += [
                "- **Trace/audit note:** display quote deterministically normalised from exact extracted substring; raw exact quote retained below and in the analyst bundle.",
                f"  - Raw exact quote: {q['exact_quote']}",
                f"  - Normalization reason: {q.get('quote_display_normalization_reason')}",
                "",
            ]
        quote_md += [f"- **Why it matters:** {q['why_it_matters']}", f"- **What it does not prove:** {q['what_it_does_not_prove']}", ""]
    (analyst_dir / "quote_bank.md").write_text("\n".join(quote_md), encoding="utf-8")
    json_dump(analyst_dir / "governance_gap_register.json", {"gaps": gaps})
    json_dump(analyst_dir / "failure_pathways.json", {"failure_pathways": pathways})
    json_dump(analyst_dir / "control_recommendations.json", {"control_recommendations": controls})
    (analyst_dir / "AI_ANALYST_PROMPT.md").write_text(build_ai_prompt(), encoding="utf-8")
    json_dump(analyst_dir / "AI_ANALYST_INPUT_BUNDLE.json", bundle)
    (analyst_dir / "AI_REPORT_VALIDATION_RULES.md").write_text(build_validation_rules(), encoding="utf-8")
    return {"quote_bank": quote_bank, "gap_register": gaps, "failure_pathways": pathways, "control_recommendations": controls, "analyst_bundle": bundle, "low_confidence_quote_candidates": low_confidence_quote_candidates}

def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Extract a local document and run the LAIF assessment report wrapper.")
    parser.add_argument("input_file", type=Path, help="Local input document path")
    parser.add_argument("--mode", choices=ASSESSMENT_MODES, default="external_framework")
    parser.add_argument("--sector", choices=SECTOR_CHOICES, default="auto")
    parser.add_argument("--source-type", default="uploaded_document")
    parser.add_argument("--document-name", default=None)
    parser.add_argument("--output-dir", type=Path, default=Path(DEFAULT_OUTPUT_DIR))
    parser.add_argument("--markdown", dest="markdown", action="store_true", default=True)
    parser.add_argument("--no-markdown", dest="markdown", action="store_false")
    parser.add_argument("--json", dest="json_output", action="store_true", default=True)
    parser.add_argument("--no-json", dest="json_output", action="store_false")
    parser.add_argument("--extractor", choices=EXTRACTOR_CHOICES, default="auto")
    parser.add_argument("--fail-on-warnings", action="store_true")
    parser.add_argument("--print-report", action="store_true")
    parser.add_argument("--no-write", action="store_true", help="Do not write markdown, JSON, or processing index outputs")
    parser.add_argument("--original-pending-path", default="", help="Original pending/source path before archival copy, for batch identity metadata")
    parser.add_argument("--stored-source-path", default="", help="Archived source path retained by batch processing, for identity metadata")
    return parser


def run(args: argparse.Namespace) -> int:
    input_path_original = str(args.input_file)
    input_path = args.input_file.expanduser().resolve()
    extraction = extract_document(input_path, args.extractor)
    if args.fail_on_warnings and extraction.warnings:
        raise ExtractionError("Extraction produced warnings and --fail-on-warnings was set: " + "; ".join(extraction.warnings))

    selected_sector = args.sector
    document_name = args.document_name or input_path.stem
    processed_at = utc_now_iso()
    source_hash = sha256_file(input_path)
    output_stem = safe_stem(input_path.stem)

    processing = build_processing_metadata(
        input_path=input_path,
        input_path_original=input_path_original,
        output_dir=args.output_dir,
        processed_at_utc=processed_at,
        source_sha256=source_hash,
        safe_output_stem=output_stem,
        markdown_enabled=args.markdown,
        json_enabled=args.json_output,
        original_pending_path=args.original_pending_path,
        stored_source_path=args.stored_source_path,
    )
    extraction_metadata = {
        "input_path_original": input_path_original,
        "input_path": str(input_path),
        "runner_input_path": str(input_path),
        "original_pending_path": args.original_pending_path or input_path_original,
        "stored_source_path": args.stored_source_path or str(input_path),
        "input_file_name": input_path.name,
        "original_file_name": input_path.name,
        "original_file_stem": input_path.stem,
        "source_sha256": source_hash,
        "safe_output_stem": output_stem,
        "extractor_requested": args.extractor,
        "extractor_used": extraction.extractor_used,
        "extraction_confidence": extraction.extraction_confidence,
        "extracted_characters": len(extraction.text),
        "minimum_extracted_characters": MIN_EXTRACTED_CHARACTERS,
        "warning_count": len(extraction.warnings),
        "warnings": extraction.warnings,
        "error_count": len(extraction.errors),
        "errors": extraction.errors,
        "ocr_performed": False,
        "network_access_used": False,
    }

    assessment = assess(
        document_name,
        args.source_type,
        extraction.text,
        sector=selected_sector,
        assessment_mode=resolve_assessment_mode(args.mode),
        source_sha256=source_hash,
        original_file_name=input_path.name,
        original_pending_path=args.original_pending_path or input_path_original,
        stored_source_path=args.stored_source_path or str(input_path),
        runner_input_path=str(input_path),
        processed_at_utc=processed_at,
    )
    base_report = generate_markdown_report([assessment])
    markdown_report = markdown_metadata_block(processing, extraction_metadata, assessment) + base_report
    payload = {
        "processing_metadata": processing,
        "extraction_metadata": extraction_metadata,
        "assessment_result": assessment,
    }

    if not args.no_write:
        args.output_dir.mkdir(parents=True, exist_ok=True)
        if args.markdown:
            Path(processing["markdown_output_path"]).write_text(markdown_report, encoding="utf-8")
        if args.json_output:
            json_dump(Path(processing["json_output_path"]), payload)
        analyst_outputs = write_institutional_outputs(args.output_dir, processing, extraction_metadata, assessment, extraction.text)
        payload["institutional_analyst_outputs"] = {
            "quote_bank_count": len(analyst_outputs["quote_bank"]),
            "gap_count": len(analyst_outputs["gap_register"]),
            "failure_pathway_count": len(analyst_outputs["failure_pathways"]),
            "control_recommendation_count": len(analyst_outputs["control_recommendations"]),
        }
        if args.json_output:
            json_dump(Path(processing["json_output_path"]), payload)
        append_index(args.output_dir, index_record(processing, extraction_metadata, assessment, input_path, document_name))

    print(f"Original input path: {input_path_original}")
    print(f"Resolved input path: {input_path}")
    print(f"Original file name: {input_path.name}")
    print(f"Processed at UTC: {processed_at}")
    print(f"Source SHA-256: {source_hash}")
    print(f"Extractor used: {extraction.extractor_used}")
    print(f"Extracted characters: {len(extraction.text)}")
    print(f"Assessment mode: {assessment.get('assessment_mode')}")
    print(f"Sector profile: {assessment.get('sector_profile')}")
    if args.no_write:
        print("Write mode: disabled (--no-write); no markdown, JSON, or index outputs written.")
    else:
        if args.markdown:
            print(f"Markdown report: {processing['markdown_output_path']}")
        if args.json_output:
            print(f"JSON report: {processing['json_output_path']}")
        print(f"Institutional report: {args.output_dir / (processing['safe_output_stem'] + '.institutional_report.md')}")
        print(f"Technical appendix: {args.output_dir / (processing['safe_output_stem'] + '.technical_appendix.md')}")
        print(f"Analyst bundle: {args.output_dir / 'analyst' / 'analyst_bundle.json'}")
        print(f"Processing index: {args.output_dir / INDEX_FILE_NAME}")
    if args.print_report:
        print("\n" + markdown_report)
    return 0


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    try:
        return run(args)
    except ExtractionError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
